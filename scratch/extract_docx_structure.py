from docx import Document
from pathlib import Path

base = Path(r"c:/Study/Uni/Sem 2/Group project/CW_2/sky-team-registry/only_for_to_read_context_no_github_push")
files = [
    base / "5COSC021W Coursework 2 - GROUP template 2024_25(1).docx",
    base / "5COSC021W Coursework 2 - INDIVIDUAL template 2024_25(1).docx",
]

for f in files:
    doc = Document(str(f))
    out = []
    out.append(f"FILE: {f.name}")
    out.append(f"PARAGRAPHS: {len(doc.paragraphs)}")
    out.append(f"TABLES: {len(doc.tables)}")
    out.append("--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            out.append(f"P{i}: {t}")
    out.append("--- TABLES ---")
    for ti, table in enumerate(doc.tables):
        out.append(f"TABLE {ti}: rows={len(table.rows)} cols~={len(table.columns)}")
        for ri, row in enumerate(table.rows):
            cells = []
            for ci, cell in enumerate(row.cells):
                txt = " ".join(cell.text.split())
                cells.append(f"C{ci}=\"{txt}\"")
            out.append(f"R{ri}: " + " | ".join(cells))
    out_path = f.with_suffix(".structure.txt")
    out_path.write_text("\n".join(out), encoding="utf-8")
    print(out_path)
