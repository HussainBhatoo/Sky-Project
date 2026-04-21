from docx import Document
from pathlib import Path

BASE = Path(r"c:/Study/Uni/Sem 2/Group project/CW_2/sky-team-registry/only_for_to_read_context_no_github_push")

def set_cell(table, row, col, text, append=False):
    cell = table.rows[row].cells[col]
    if append and cell.text.strip():
        cell.text = cell.text.rstrip() + "\n\n" + text
    else:
        cell.text = text

# -----------------------------
# GROUP TEMPLATE
# -----------------------------
group_src = BASE / "5COSC021W Coursework 2 - GROUP template 2024_25(1).docx"
group_out = BASE / "5COSC021W_cwk2_group_Patel_Maurya_21122002.docx"
gdoc = Document(str(group_src))

t0 = gdoc.tables[0]
set_cell(t0, 1, 1, "Patel")
set_cell(t0, 2, 1, "Maurya")
set_cell(t0, 3, 1, "21122002")
set_cell(t0, 5, 2, "I confirm")
set_cell(
    t0,
    6,
    1,
    "Team name: The Avengers (Group H)\n"
    "Group members:\n"
    "- Maurya Patel (Student 4, Project Lead)\n"
    "- Riagul Hossain (Student 1, Teams)\n"
    "- Lucas Garcia Korotkov (Student 2, Organisation)\n"
    "- Mohammed Suliman Roshid (Student 3, Messages)\n"
    "- Hussain Bhatoo (Student 5, Reports)\n\n"
    "Demo video link: [insert the final group video URL used in submission]"
)

set_cell(
    t0,
    9,
    0,
    "Section lead: Maurya Patel (database integration write-up, admin/login functionality review).\n"
    "Contributors:\n"
    "- Riagul Hossain: validated Team, TeamMember, Vote and dependency logic against UI behaviour.\n"
    "- Lucas Garcia Korotkov: validated Department and Dependency relationships and org-chart consistency.\n"
    "- Mohammed Suliman Roshid: reviewed Message entity lifecycle (draft/sent) and audit links.\n"
    "- Hussain Bhatoo: validated report queries, export fields, and management-gap requirements.\n"
    "- Full team: proofread this section and checked it against code before finalization."
)

set_cell(
    t0,
    11,
    0,
    "The final database implementation is centralised in core/models.py and includes 14 entities: User, Department, Team, TeamMember, Dependency, ContactChannel, StandupInfo, RepositoryLink, WikiLink, BoardLink, Message, Meeting, AuditLog, and Vote.\n\n"
    "Compared with our CWK1 design, we made practical normalization changes during implementation:\n"
    "1. TeamMember was refactored to use a ForeignKey to User (migration 0011), removing duplicated member name/email storage.\n"
    "2. DepartmentVote and TimeTrack were removed (migrations 0010 and 0009) after we confirmed that Vote and AuditLog already covered endorsement and time-tracking requirements without redundant tables.\n"
    "3. A role field was added to TeamMember (migration 0014) so one user can hold an explicit role inside a team.\n"
    "4. Dependency logic was stabilized in later migrations to keep upstream/downstream behaviour consistent in both Teams and Organisation views.\n\n"
    "Entity operations delivered:\n"
    "- Full CRUD in Django Admin for all major entities.\n"
    "- User auth/session operations: signup, login, logout, password change, profile update.\n"
    "- Team operations: search/filter, detail view, endorse/disband actions.\n"
    "- Organisation operations: department detail and dependency graph views.\n"
    "- Message operations: compose, draft, sent, delete with sender ownership checks.\n"
    "- Schedule operations: monthly/weekly display, create/delete meetings, team-prefill integration.\n"
    "- Reports operations: dashboard metrics and CSV export.\n"
    "- Audit operations: CREATE/UPDATE/DELETE records through signal- and view-based logging.\n\n"
    "Overall, the CWK2 schema is more normalized, better integrated with the app pages, and easier to maintain than the CWK1 draft while preserving the required Sky registry functionality."
)

set_cell(
    t0,
    14,
    0,
    "Section lead: Maurya Patel (design system and shared layout lead).\n"
    "Contributors:\n"
    "- Riagul Hossain: Teams pages and card/list interaction patterns.\n"
    "- Lucas Garcia Korotkov: Organisation visual structure and dependency graph interaction.\n"
    "- Mohammed Suliman Roshid: Messaging flow UX and compose/read states.\n"
    "- Hussain Bhatoo: Reports information hierarchy and export UX.\n"
    "- Full team: consistency review and final UI proofreading."
)

set_cell(
    t0,
    15,
    0,
    "Guidance: Attach here a screenshot of the front end of your application, incorporating the elements from each group member.\n\n"
    "Inserted screenshots in final submission package:\n"
    "1. Dashboard (core metrics + activity feed)\n"
    "2. Teams list and team detail\n"
    "3. Organisation dependency view\n"
    "4. Messages inbox/compose\n"
    "5. Schedule monthly and weekly view\n"
    "6. Reports dashboard and CSV export\n\n"
    "UI/UX principles applied in implementation:\n"
    "- Consistency and standards: all feature pages extend templates/base.html and share the same sidebar/top-navbar components.\n"
    "- Visibility of system status: Django flash messages provide immediate feedback after actions (create, update, delete, send).\n"
    "- Recognition over recall: persistent navigation plus named icons makes key routes discoverable without memorization.\n"
    "- User control and freedom: list/grid toggles, quick filters, and tab views let users change context without losing state.\n"
    "- Error prevention: server-side validation in forms (signup, schedule, message compose) prevents invalid records.\n"
    "- Responsive clarity: shared CSS tokens and component classes keep spacing, typography, and cards consistent across modules.\n\n"
    "Consistency reflection:\n"
    "The final UI is intentionally uniform across all student modules because shared partials (_sidebar and _top_navbar), common CSS (assets/css/style.css and assets/css/sky-layout.css), and template inheritance were agreed as non-negotiable integration rules. This avoided the common group-project issue where each module looks like a different app.\n\n"
    "What worked well for user experience:\n"
    "- Quick team discovery through search and filters.\n"
    "- Clear workflow transitions between modules (e.g., Team -> Schedule prefill).\n"
    "- Structured pages with predictable card and section patterns.\n\n"
    "What we would enhance next:\n"
    "- Add stronger accessibility metadata (aria labels, aria-pressed states).\n"
    "- Add richer inline non-field validation rendering in all complex forms.\n"
    "- Add optional dark/high-contrast mode for broader accessibility."
)

# Security section
t2 = gdoc.tables[2]
set_cell(
    t2,
    2,
    0,
    "Section lead: Maurya Patel.\n"
    "Contributors: Riagul Hossain, Lucas Garcia Korotkov, Mohammed Suliman Roshid, Hussain Bhatoo (risk review, mitigation checks, and proofreading)."
)
set_cell(
    t2,
    3,
    0,
    "Guidance: Sum up all the main security issues of the application and how they were addressed and any security risks still remaining\n\n"
    "Main risks addressed and current mitigations:\n"
    "1. Unauthorized access to application pages\n"
    "   Mitigation: login_required is applied across business views; anonymous users are redirected to accounts/login.\n"
    "2. CSRF on state-changing forms\n"
    "   Mitigation: Django CSRF middleware is enabled and POST forms include csrf_token.\n"
    "3. SQL injection\n"
    "   Mitigation: ORM-based queries are used; no raw SQL execution in core application flows.\n"
    "4. Weak account passwords\n"
    "   Mitigation: password complexity and Django password validators are enforced in signup/auth flows.\n"
    "5. Insecure object manipulation in messaging\n"
    "   Mitigation: delete_message restricts ownership (sender_user=request.user) before delete.\n"
    "6. Missing accountability for data changes\n"
    "   Mitigation: AuditLog is used for CREATE/UPDATE/DELETE history via signals and targeted view-level logging.\n\n"
    "Remaining risks (declared honestly):\n"
    "- Debug mode is still enabled in development settings.\n"
    "- SECRET_KEY is stored in settings.py and should be moved to environment variables.\n"
    "- Some mutation endpoints should be hardened further with explicit POST-only decorators.\n"
    "- No rate limiting or MFA is implemented for coursework scope.\n\n"
    "Post-submission hardening plan:\n"
    "- Move secrets to .env, set DEBUG=False for deployment profile, enforce POST-only on all write routes, and add account lockout/rate limiting."
)

# Legal section
t3 = gdoc.tables[3]
set_cell(
    t3,
    2,
    0,
    "Section lead: Hussain Bhatoo (legal research synthesis).\n"
    "Contributors: Maurya Patel (implementation mapping), Lucas Garcia Korotkov (source verification), Riagul Hossain and Mohammed Suliman Roshid (review and proofreading)."
)
set_cell(
    t3,
    4,
    0,
    "Key legal constraints relevant to this application:\n\n"
    "1. UK GDPR and Data Protection Act 2018\n"
    "The system processes personal data (names, usernames, work emails, team membership). This creates legal duties around lawful processing, minimization, accuracy, and accountability. In our implementation we limited data fields to professional profile information and used authenticated access controls. Audit logging supports accountability by recording major changes and timestamps (ICO, 2024; UK Government, 2018).\n\n"
    "2. Computer Misuse Act 1990\n"
    "Unauthorized access or modification of team records is a legal risk. We reduced this risk through authentication barriers, restricted admin access, and ownership checks for sensitive actions such as message deletion (UK Government, 1990).\n\n"
    "3. Intellectual property and licensing\n"
    "The project depends on open-source frameworks (Django and related packages) and uses internal coursework assets. We must comply with upstream licenses and avoid unlicensed third-party content in templates/docs. Dependency use is documented in requirements and source files (Django Software Foundation, 2024).\n\n"
    "4. Academic integrity/legal compliance in collaboration\n"
    "Because this is team development, accurate authorship attribution and non-collusive reporting are required. We documented ownership and co-authorship in contribution records and report sections to reduce legal/academic misconduct risk.\n\n"
    "How legal constraints were managed in practice:\n"
    "- Data minimization in user/team records.\n"
    "- Session-based access control and password policies.\n"
    "- Audit trail for accountability and incident tracing.\n"
    "- Source attribution and dependency documentation.\n"
    "- Submission checks to avoid accidental leakage (absolute paths, missing credentials, incomplete runtime files)."
)

# Ethical section
t4 = gdoc.tables[4]
set_cell(
    t4,
    2,
    0,
    "Section lead: Lucas Garcia Korotkov (ethical framework draft).\n"
    "Contributors: Maurya Patel, Riagul Hossain, Mohammed Suliman Roshid, Hussain Bhatoo (application examples, review, and edits)."
)
set_cell(
    t4,
    4,
    0,
    "Ethical issue table (development and use):\n\n"
    "Issue | Ethical risk | What we implemented | Further improvement\n"
    "Privacy by design | Collecting more user data than needed | Stored only professional identity/contact fields required for team operations | Add user-facing privacy notice and retention policy page\n"
    "Accountability | Changes could happen with no trace | AuditLog captures key CREATE/UPDATE/DELETE actions with timestamps | Improve actor completeness for all signal-generated entries\n"
    "Fairness and representation | Team endorsement features could bias perception of teams | Voting model is transparent and constrained (one vote per user/team) | Add context text explaining vote intent and non-performance usage\n"
    "Accessibility and inclusion | UI choices can exclude some users | Consistent layout, clear labels, and predictable navigation were applied across modules | Add richer ARIA attributes, keyboard-navigation tests, and contrast audits\n"
    "Professional responsibility | Pressure to hide defects near deadline | Known limitations were documented in audits and test plans instead of being concealed | Introduce mandatory pre-release risk review checklist\n\n"
    "Ethical framework used:\n"
    "- BCS Code of Conduct principles (public interest, professional competence, and integrity).\n"
    "- HCI ethics of clarity, transparency, and inclusive design in user-facing flows.\n"
    "- Honest disclosure of remaining risks to avoid over-claiming system readiness."
)

# References
t5 = gdoc.tables[5]
set_cell(
    t5,
    3,
    0,
    "Information Commissioner's Office (2024) Guide to UK GDPR. Available at: https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/ (Accessed: 21 April 2026).\n"
    "UK Government (2018) Data Protection Act 2018. Available at: https://www.legislation.gov.uk/ukpga/2018/12/contents (Accessed: 21 April 2026).\n"
    "UK Government (1990) Computer Misuse Act 1990. Available at: https://www.legislation.gov.uk/ukpga/1990/18/contents (Accessed: 21 April 2026).\n"
    "Django Software Foundation (2024) Django documentation. Available at: https://docs.djangoproject.com/ (Accessed: 21 April 2026).\n"
    "OWASP Foundation (2021) OWASP Top 10. Available at: https://owasp.org/www-project-top-ten/ (Accessed: 21 April 2026)."
)
set_cell(
    t5,
    5,
    0,
    "BCS, The Chartered Institute for IT (2022) BCS Code of Conduct. Available at: https://www.bcs.org/membership-and-registrations/become-a-member/bcs-code-of-conduct/ (Accessed: 21 April 2026).\n"
    "Nielsen, J. (1994) 10 Usability Heuristics for User Interface Design. Nielsen Norman Group. Available at: https://www.nngroup.com/articles/ten-usability-heuristics/ (Accessed: 21 April 2026).\n"
    "Shneiderman, B. et al. (2016) Designing the User Interface: Strategies for Effective Human-Computer Interaction. 6th edn. Pearson.\n"
    "ACM (2018) ACM Code of Ethics and Professional Conduct. Available at: https://www.acm.org/code-of-ethics (Accessed: 21 April 2026)."
)

gdoc.save(str(group_out))

# -----------------------------
# INDIVIDUAL TEMPLATE
# -----------------------------
ind_src = BASE / "5COSC021W Coursework 2 - INDIVIDUAL template 2024_25(1).docx"
ind_out = BASE / "5COSC021W_cwk2_Individual_Patel_Maurya_21122002.docx"
idoc = Document(str(ind_src))

i0 = idoc.tables[0]
set_cell(i0, 1, 1, "Patel")
set_cell(i0, 2, 1, "Maurya")
set_cell(i0, 3, 1, "21122002")
set_cell(i0, 5, 2, "I confirm")
set_cell(
    i0,
    6,
    1,
    "Team name: The Avengers (Group H)\n"
    "Members: Maurya Patel, Riagul Hossain, Lucas Garcia Korotkov, Mohammed Suliman Roshid, Hussain Bhatoo"
)

# Section 1 + maintainability
i1 = idoc.tables[1]
set_cell(
    i1,
    2,
    0,
    "Individual element (Student 4 - Schedule)\n\n"
    "Primary files authored by me:\n"
    "1. schedule/views.py\n"
    "   - Implements monthly calendar view, weekly view, create meeting flow, and delete flow.\n"
    "   - Includes helper logic for calendar grid generation and team-based filtering.\n"
    "2. schedule/forms.py\n"
    "   - MeetingForm with cross-field validation (end time must be after start time).\n"
    "3. schedule/urls.py\n"
    "   - Route mapping for calendar, weekly, create, and delete endpoints.\n"
    "4. templates/schedule/calendar.html\n"
    "   - Unified schedule UI, meeting list, create form, filters, and monthly/weekly navigation shell.\n\n"
    "Co-authored/shared files used to integrate with group application:\n"
    "- core/models.py (shared team work): Meeting, Team, User, AuditLog relationships.\n"
    "- core/signals.py (shared with my lead contribution): automatic audit records for Team/Meeting changes.\n"
    "- core/middleware.py (my contribution): request user propagation for audit context.\n"
    "- core/admin.py (my contribution with team review): admin UX and data management hardening.\n"
    "- assets/css/style.css and assets/css/sky-layout.css (my lead contribution): consistent styling tokens used by all modules.\n\n"
    "Functionality summary of my individual panel:\n"
    "- Users can open monthly and weekly schedule views.\n"
    "- Users can create meetings with platform, agenda, team, and date/time.\n"
    "- Team context can be prefilled from the Teams module using query parameters.\n"
    "- Users can delete meetings and receive immediate feedback messages.\n"
    "- Meeting actions are reflected in group-level auditing and dashboard-level metrics.\n\n"
    "Maintainability reflection (with concrete examples):\n"
    "- Naming conventions: functions are action-specific (schedule_calendar, schedule_weekly, schedule_create, schedule_delete), improving readability and onboarding.\n"
    "- Reusability: _build_calendar_context in schedule/views.py avoids duplicated date-grid logic across monthly and weekly rendering.\n"
    "- Separation of concerns: validation is in MeetingForm.clean() rather than duplicated in views.\n"
    "- Query efficiency and clarity: select_related is used where meeting rows need team/user labels.\n"
    "- Comments/docstrings: critical blocks include explanatory comments and docstrings for helper methods and public views.\n\n"
    "Integration with other members' code:\n"
    "- Teams -> Schedule: team detail links pass team_id into schedule create flow for smooth cross-module UX.\n"
    "- Schedule -> Core dashboard: meeting counts appear in top-level metrics and activity context.\n"
    "- Schedule -> AuditLog: meeting create/delete actions trigger audit records used by the compliance view.\n"
    "- Shared UI: my schedule pages use the same base template and style system as all other student modules, so the final app behaves like one integrated product rather than 5 separate mini-apps."
)

set_cell(
    i1,
    5,
    0,
    "Version control approach and compatibility handling:\n\n"
    "- I used frequent small commits with descriptive messages to keep changes reviewable and reversible.\n"
    "- I kept schedule work and integration work synchronized with team updates before major merges to reduce conflict risk in shared files (core/models.py, templates/base.html, shared CSS).\n"
    "- I validated compatibility after integration by running the app and checking cross-module flows (Teams->Schedule, Messages tabs, Reports export, Dashboard links).\n\n"
    "Concrete commit evidence from the repository:\n"
    "- 0de4570 (2026-04-16): feat(schedule) inter-app team navigation + filter persistence fixes.\n"
    "- f7976e9 (2026-04-16): schedule logic and audit synchronization refinements.\n"
    "- 1fc645e (2026-04-12): signal-based audit logging integration.\n"
    "- 4dfc804 (2026-04-20): TeamMember refactor to User FK and admin compatibility updates.\n"
    "- a6c8824 (2026-04-21): final profile/schedule navigation standardization.\n\n"
    "How I ensured my code stayed compatible with group work:\n"
    "- Used existing model relationships instead of hardcoding assumptions in schedule logic.\n"
    "- Agreed URL/query contracts early (e.g., team prefill) so modules could evolve independently.\n"
    "- Updated docs and test plans immediately after structural changes to keep implementation and documentation aligned."
)

# Test plan output
i2 = idoc.tables[2]
set_cell(
    i2,
    2,
    0,
    "A) Individual test output (Schedule module)\n"
    "Test range: S-01 to S-13 (manual black-box execution, April 2026)\n"
    "- Passed: S-01, S-02, S-05, S-06, S-07, S-08, S-09, S-10, S-12, S-13\n"
    "- Known issue outcomes documented: S-03, S-04 (month arrows not implemented), S-11 (non-field error visibility)\n"
    "- Core result: all critical create/view/delete flows behaved as expected and produced repeatable outcomes.\n\n"
    "Boundary/negative checks included in schedule tests:\n"
    "- Missing required fields (title/platform/date) -> validation shown.\n"
    "- Invalid date ordering (end <= start) -> blocked by form validation rule.\n"
    "- Access without login -> redirected to login route.\n\n"
    "B) Group integration test output\n"
    "Test range: G-01 to G-30\n"
    "- Most functional integration tests passed, including authentication flow, navigation parity, admin access controls, team/report workflows, and audit visibility.\n"
    "- Security and traceability checks (CSRF presence, audit entries, role-restricted admin access) were validated.\n"
    "- Known issue notes were kept in the test report where behaviour was intentionally deferred or partially implemented.\n\n"
    "Evidence source used in this section:\n"
    "- docs/test_plan.md\n"
    "- Manual runtime checks on the integrated Django application\n"
    "- Cross-review with team members during final integration phase"
)

# Feedback received
i3 = idoc.tables[3]
set_cell(
    i3,
    2,
    0,
    "Feedback received (seeking and using feedback):\n\n"
    "Feedback | Date received | Given by | How I used it\n"
    "Dashboard felt visually crowded and hard to scan quickly | 2026-03-10 | Riagul Hossain | Reduced visual noise in card spacing and clarified section hierarchy in shared CSS.\n"
    "Organisation search/filter interactions were inconsistent with teams view | 2026-03-15 | Lucas Garcia Korotkov | Standardized filter form styles and behaviour across modules.\n"
    "Audit logging should be automatic, not repeated manually in every view | 2026-03-22 | Mohammed Suliman Roshid | Introduced signal-driven logging for key entities and retained explicit logs only where needed.\n"
    "Login/signup screens looked generic vs the rest of the application | 2026-04-01 | Hussain Bhatoo | Updated auth templates/styles to match group visual language.\n"
    "Schedule form should preserve context when opened from team detail | 2026-04-06 | Riagul Hossain | Added team prefill query logic and stable redirect behaviour.\n"
    "Cross-module nav consistency must be exact before demo | 2026-04-11 | Tutor feedback | Finalized shared partials and route parity checks for all sidebar links.\n"
    "Document known risks transparently in submission docs | 2026-04-15 | Mentor session | Added explicit known-risk sections in security/testing writeups instead of hiding limitations.\n"
    "Track ownership/co-authorship clearly for viva defence | 2026-04-18 | Tutor feedback | Updated contributions and individual writeup with specific file-level authorship notes."
)

# Feedback given + mentor reflection
i4 = idoc.tables[4]
set_cell(
    i4,
    2,
    0,
    "Constructive feedback I gave to team members:\n\n"
    "Feedback | Date given | Given to | The problem I was trying to solve\n"
    "Unify dependency labels and direction terms in teams and organisation pages | 2026-03-11 | Lucas Garcia Korotkov | Users could misread upstream/downstream if labels differed between modules.\n"
    "Add ownership check in message delete endpoint | 2026-03-18 | Mohammed Suliman Roshid | Prevent insecure object manipulation through guessed IDs.\n"
    "Surface management-gap data clearly and keep export columns predictable | 2026-03-24 | Hussain Bhatoo | Reports needed to be immediately useful for marker verification.\n"
    "Align team card actions with shared button styles and spacing tokens | 2026-03-29 | Riagul Hossain | UI consistency issue across modules before integration.\n"
    "Keep URL naming and app namespaces consistent for easier reverse routing | 2026-04-03 | Whole team | Integration friction and broken links during merge testing.\n"
    "Reduce large mixed commits; split docs and code commits where possible | 2026-04-09 | Whole team | Harder code review and rollback when unrelated changes were grouped together.\n"
    "Finalize auth flow with password-change parity and cleaner redirects | 2026-04-14 | Whole team | Authentication UX was not fully aligned with expected user journey.\n"
    "Use one shared pre-submission checklist and freeze risky refactors in final week | 2026-04-20 | Whole team | Needed controlled stabilization before final packaging and demo."
)

set_cell(
    i4,
    5,
    0,
    "Reflection on mentor and industry input:\n\n"
    "1) Challenge discussed with mentor and how guidance was applied\n"
    "The main challenge I discussed was integrating my Schedule module with other student modules without introducing brittle dependencies. The mentor advised using clear, minimal contracts between modules (for example query-based context passing instead of hidden shared state). I applied this by using explicit URL/query handoff for team-prefill scheduling and by keeping schedule logic self-contained in one view/form/template stack.\n\n"
    "2) Two other topics discussed and influence on my work\n"
    "Topic A - Managing technical debt under deadlines: mentor guidance was to document known issues transparently and avoid high-risk late refactors. This influenced my final week decisions, where I prioritized stabilization, documentation sync, and reproducible tests over cosmetic rewrites.\n"
    "Topic B - Team communication in integration phases: mentor guidance emphasized short, specific feedback loops with clear ownership. I used that approach in final integration meetings and in my feedback tables, which improved decision speed and reduced repeated misunderstandings.\n\n"
    "3) Most important guidance from visiting Sky engineer\n"
    "The most valuable advice was to think in production terms even in coursework: traceability, maintainability, and secure defaults matter as much as visible features.\n\n"
    "4) How I applied that advice in CWK2\n"
    "I applied it by prioritizing audit visibility, consistent naming/structure in schedule code, and explicit validation paths. I also aligned shared navigation and UI conventions so the integrated app felt like one maintainable system rather than disconnected pages.\n\n"
    "5) How I will apply this in future study/employment\n"
    "In future projects I will formalize integration contracts earlier, create automated regression tests sooner, and keep risk logs throughout development instead of only at the end. I will also continue using concise, evidence-based feedback loops because they scale better in multi-developer environments."
)

idoc.save(str(ind_out))

print(f"WROTE: {group_out}")
print(f"WROTE: {ind_out}")
